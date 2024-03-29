import sys
from time import sleep
from datetime import datetime
from tempfile import gettempdir
from re import compile
from os.path import isfile, join as pathJoin
from os import chdir, environ, stat, remove
from csv import DictReader
import json


def excepthook(type, value, tb):
    import traceback
    traceback.print_exception(type, value, tb)
    input()


sys.excepthook = excepthook

try:
    from requests import get
    from docx import Document
except:
    from subprocess import check_call
    from sys import executable
    for package in ["requests", "python-docx"]:
        check_call([executable, "-m", "pip", "install", package])
    from requests import get
    from docx import Document

# environ["PYTHONIOENCODING"] = "utf-8"

temp_dir = gettempdir()

# chdir(temp_dir)

now = datetime.now().timestamp()


def get_csv(url):
    filename = pathJoin(temp_dir, url[url.rfind("/")+1:])
    if not isfile(filename) or now - stat(filename).st_mtime > 60 * 60:
        print("downloading " + url)
        file_csv = get(url)
        # print(str(len(file_csv.text)) + " downloaded to " + filename)
        open(filename, "wt").write(file_csv.text)
    return DictReader(open(filename, "rt"))


def get_json(url, filename=False):
    if not filename:
        filename = url[url.rfind("/")+1:]
    filename = pathJoin(temp_dir, filename)
    if not isfile(filename) or now - stat(filename).st_mtime > 60 * 60:
        print("downloading " + url)
        file_json = get(url)
        # print(str(len(file_csv.text)) + " downloaded to " + filename)
        open(filename, "wt").write(file_json.text)
    return json.load(open(filename, "rt"))["value"]


def get_laws():
    # return get_csv("https://production.oknesset.org/pipelines/data/bills/kns_bill/kns_bill.csv")
    return get_json("https://knesset.gov.il/OdataV4/ParliamentInfo/KNS_Bill", "KNS_Bill.json")


def get_docs():
    # return get_csv("https://production.oknesset.org/pipelines/data/bills/kns_documentbill/kns_documentbill.csv")
    return get_json("https://knesset.gov.il/OdataV4/ParliamentInfo/KNS_DocumentBill", "KNS_DocumentBill.json")


def get_doc(url, retry=3):
    filename = pathJoin(temp_dir, url[url.rfind("/")+1:])
    if(not isfile(filename)):
        print("downloading from " + url + " to " + filename)
        data = get(url)  # , verify=False)
        #print(str(len(data.text)) + " downloaded from " + url + " to " + filename)
        #data.raw.decode_content = True
        open(filename, "wb").write(data.content)
    try:
        return Document(filename)
    except Exception as e:
        if retry > 0:
            print(filename + " failed to open, retrying")
            remove(filename)
            sleep(500)
            return get_doc(url, retry=retry-1)
        else:
            raise e


laws = {}
for law in get_laws():
    laws[law["Id"]] = law

print(str(len(laws)) + " laws loaded")

dups = compile("(פ[/\\\\][0-9]+[/\\\\]2(1|3|4|5|2))")
init = compile("יוז(מות|מים|מת|ם):\t? +חבר(ות|י|ת)? הכנסת\t")
numbers = compile("[\d]+")
scored_laws = {}
for name in ["laws21", "laws22", "laws23", "laws24"]:
    dict = DictReader(open(name + ".csv", "rt"))
    for line in dict:
        if line.get("מספר חוק") and line.get("ניקוד לחוק") != None and line.get("ניקוד לחוק") != "":
            scored_laws[line["מספר חוק"]] = line

scores = [['"שם הצעת החוק","מדרג","מספר חוק","ניקוד לחוק", "קישור להצעה", "הסבר הדירוג","הערות אחרות","הגיע להצבעה?","עבר?","יוזם ראשון","חתומים"']] + [[]] * 5000
n = 1
CURRENT_KNESSET = "25"
for line in DictReader(open("laws" + CURRENT_KNESSET + ".csv", "rt")):
    if line.get("מספר חוק") and not line.get("מספר חוק") in [("פ/" + str(n) + "/" + CURRENT_KNESSET), ("פ\\" + str(n) + "\\" + CURRENT_KNESSET), ("פ\\" + CURRENT_KNESSET + "\\" + str(n))]:
        n += 1
        if not line.get("מספר חוק") in [("פ/" + str(n) + "/" + CURRENT_KNESSET), ("פ\\" + str(n) + "\\" + CURRENT_KNESSET), ("פ\\" + CURRENT_KNESSET + "\\" + str(n))]:
            n -= 2
            if not line.get("מספר חוק") in [("פ/" + str(n) + "/" + CURRENT_KNESSET), ("פ\\" + str(n) + "\\" + CURRENT_KNESSET), ("פ\\" + CURRENT_KNESSET + "\\" + str(n))]:
                n += 1
                print(line.get("מספר חוק"), n, "no match")

    if not line.get("מספר חוק"):
        line["מספר חוק"] = ("פ/" + str(n) + "/" + CURRENT_KNESSET)
    if line.get("הסבר הדירוג").find("ראה חוק") > 0 and not line.get("מדרג"):
        line["מדרג"] = "dup_laws_bot"
    # print(line)
    scores[n] = [
        "\"" + line.get("שם הצעת החוק") +
        "\"", line.get("מדרג"), line.get("מספר חוק"),
        line.get("ניקוד לחוק") or line.get("ניקוד"), line.get(
            "קישור להצעה"), "\""+line.get("הסבר הדירוג").replace("\"", "'")+"\"",
        line.get("הערות אחרות"), line.get("הגיע להצבעה?"), line.get("עבר?"), line.get("יוזם ראשון"), line.get("חתומים")]
    if line.get("ניקוד לחוק") != None and line.get("ניקוד לחוק") != "":
        if line.get("מספר חוק"):
            scored_laws[line["מספר חוק"]] = line
    n += 1
# sys.exit(0)
duplicates = {}
laws_initiators = [[]]*5000

news_csv = 'קישור, שם, מספר\n'
unscored_csv = 'קישור, שם, מספר, עלה להצבעה\n'
old_csv = 'קישור, שם, מספר חדש, דירוג, מספר קודם\n'

split_initiators = compile("[\n\t]")
i = 0
laws_last = 0
knst = {}
DEBUG_STATUSES = False
for doc in get_docs():
    i += 1
    # if i%1000 == 0:
    # 	print(i)

    law = laws.get(doc["BillID"])
    if not law:
        print("law " + str(doc["BillID"]) + " not found")
        # print(doc)
        continue

    if (not DEBUG_STATUSES) and str(law["KnessetNum"]) != CURRENT_KNESSET:
        knst[law["KnessetNum"]] = True
        continue

    # 53 - הצעה ממשלתית
    # 54 - הצעה פרטית
    # 55 - ועדה
    # 6042 - נוסח חדש
    # 6041 - נוסח משולב
    # 6043 - מנדטורי
    # 6045 - מועצת המדינה הזמנית
    if str(law["SubTypeID"]) != '54':
        if DEBUG_STATUSES and str(law["SubTypeID"]) not in ["53", "54", "55", "6042", "6045", "6043", "6041"]:
            print("unknown law type", law["SubTypeID"],
                  "-", law["SubTypeDesc"], law["Id"])
        else:
            continue

    # 0 - מסמכים לא משויכים
    # 1 - דיון מוקדם
    # 2 - הצעת חוק לקריאה הראשונה
    # 3 - הצעת חוק לקריאה הראשונה - נוסח מתוקן
    # 4 - הצעת חוק לקריאה השנייה והשלישית
    # 5 - הצעת חוק לקריאה השנייה והשלישית - לוח תיקונים
    # 7 - הצעת חוק לקריאה השלישית - לוח תיקונים
    # 8 - חוק - נוסח לא רשמי
    # 9 - חוק - פרסום ברשומות
    # 12 - מסמך מ.מ.מ
    # 17 - החלטת ממשלה
    # 23 - פרוטוקול ועדה
    # 45 - קטע מדברי הכנסת
    # 28 - דברי הכנסת
    # 46 - הצעת חוק לקריאה השנייה והשלישית - הנחה מחדש
    # 47 - הצעת חוק לקריאה השנייה והשלישית - הנחה מחדש
    # 49 - הצעת חוק לקריאה השנייה והשלישית - הנחה מחדש
    # 50 - הצעת חוק לקריאה השנייה והשלישית - לוח תיקונים
    # 51 - הצעת חוק לדיון מוקדם - נוסח מתוקן
    # 56 - הצעת חוק לקריאה הראשונה - נוסח לדיון בוועדה
    # 57 - תיקון טעות בחוק שהתקבל
    # 58 - הודעה לעיתונות
    # 59 - חומר רקע
    # 60 - הצעת חוק לקריאה השניה והשלישית - נוסח לדיון בוועדה
    # 97 - חוק - נוסח חדש
    # 98 - הצעת נוסח חדש
    # 99 - חוק - תיקון טעות
    # 101 - הצעת חוק לקריאה השניה והשלישית - פונצ  בננה
    # 102 - הצעת חוק לקריאה השניה והשלישית - לוח תיקונים - פונצ בננה
    # 103 - הצעת חוק לקריאה השנייה והשלישית - הנחה מחדש- פונצ בננה
    # 104 - הצעת חוק לקריאה השנייה והשלישית - הנחה מחדש-פונצ בננה
    # 105 - הצעת חוק לקריאה השנייה והשלישית - לוח תיקונים-פונצ בננה
    # 118 - חוק - תיקון טעות - פרסום ברשומות
    # 122 - הערכת השפעות רגולציה
    # 125 - חוק - פרסום ברשומות
    # 127 - הצעת חוק - הודעת מערכת רשומות
    if str(doc["GroupTypeID"]) != '1':
        if str(doc["GroupTypeID"]) not in [
            '0', '1', '2', '3', '4', '5', '7', '8', '9', '12',
            '15', '17', '23', '28', '45', '46', '47', '49',
            '50', '51', '56', '57', '58', '59', '60', '97',
            '98', '99', '101', '102', '103', '104', '105',
            '118', '122', '125', '127'
        ]:
            print("unknown doc type",
                  doc["GroupTypeID"], "-", doc["GroupTypeDesc"])
        continue

    num = 1 if DEBUG_STATUSES else int(law["PrivateNumber"])
    laws_last += 1
    # if num != laws_last:
    # 	print(num, laws_last)
    law_name = "פ\\{}\\{}".format(CURRENT_KNESSET, num)
    if len(scores) <= num or not scores[num]:
        scores[num] = ["\"" + law["Name"].replace("\"", "'") + "\""] + [''] * 8
    if not scores[num][2]:
        scores[num][2] = law_name
    if not scores[num][4]:
        scores[num][4] = 'https://main.knesset.gov.il/Activity/Legislation/Laws/Pages/LawBill.aspx?t=lawsuggestionssearch&lawitemid=' + \
            str(law["Id"])
    # print(law)

    # StatusID
    # 104 - הונחה על שולחן הכנסת לדיון מוקדם
    # 106 - בוועדת הכנסת לקביעת הוועדה המטפלת
    # 108 - בהכנה לקריאה הראשונה בוועדה
    # 109 - אושרה בוועדה לקריאה ראשונה
    # 111 - לדיון במליאה לקראת הקריאה הראשונה
    # 113 - בהכנה לקריאה שנייה-שלישית בוועדה
    # 114 - לדיון במליאה לקראת קריאה שנייה-שלישית
    # 115 - הוחזרה לוועדה להכנה לקריאה שלישית
    # 116 - בהכנה לקריאה שנייה-שלישית בוועדה
    # 118 - חוק עבר
    # 120 - לדיון במליאה על החלת דין רציפות
    # 122 - מוזגה עם הצעת חוק אחרת
    # 124 - הוסבה להצעה לסדר היום
    # 126 - לאישור מיזוג בוועדת הכנסת
    # 130 - הונחה על שולחן הכנסת לקריאה שנייה-שלישית
    # 140 - להסרה מסדר היום לבקשת ועדה
    # 141 - הונחה על שולחן הכנסת לקריאה ראשונה
    # 142 - בוועדת הכנסת לקביעת הוועדה המטפלת
    # 143 - להסרה מסדר היום לבקשת ועדה
    # 150 - במליאה לדיון מוקדם
    # 161 - לאישור פיצול במליאה
    # 167 - אושרה בוועדה לקריאה ראשונה
    # 169 - לאישור מיזוג בוועדת הכנסת
    # 175 - בדיון בוועדה על החלת דין רציפות
    # 177 - החקיקה נעצרה
    # 178 - אושרה בוועדה לקריאה שנייה-שלישית

    # PostponementReasonID
    # 41 - כהונת חה"כ פסקה בשל התחדשות חברות ח"כ אחר
    # 1065
    # 2245
    # 2506 הסרה מסד"י בהמלצת ועדה אחרי ד. מוקדם
    # 2507 - לא נתקבלה בק-1
    # 2508 - הסרה מסד"י בהמלצת ועדה אחרי ק-1
    # 2509 - לא נתקבלה בק-2
    # 2510 - לא נתקבלה בק-3
    # 2511 - חזרת חה"כ המציע לפני ד. מוקדם
    # 2512 - חזרת חה"כ המציע אחרי ד. מוקדם
    # 2505 - הסרה מסד"י בד. מוקדם
    # 3010 - מונה לנשיא
    # 3011 - חה"כ המציע נפטר
    # 3012 - חה"כ המציע התפטר
    # 3013 - חה"כ המציע מונה לתפקיד בממשלה
    # 3087 - הצעת החוק לא נדונה/לא הוצבעה במועד
    # 3086 - לא הושג הרוב הדרוש - הצ"ח תקציבית
    # 3112 - לא הושג הרוב הדרוש

    if str(law["StatusID"]) in ["104"]:
        scores[num][7] = "0"
        scores[num][8] = "0"
    elif str(law["StatusID"]) in ["118"]:
        scores[num][7] = "1"
        scores[num][8] = "1"
    elif str(law["StatusID"]) == "177":
        if(str(law["PostponementReasonID"]) in [
            '41', '1065', '2511', '2512', '3010',
            '3011', '3012', '3013', '3087'
        ]):
            scores[num][7] = "0"
            scores[num][8] = "0"
        else:
            if not str(law["PostponementReasonID"]) in [
                '2505', '2506', '2507', '2508', '2509',
                '2510', '3086', '3087', '3112'
            ]:
                print("unknown PostponementReason", law["PostponementReasonID"],
                      law["PostponementReasonDesc"], law["Id"], law["KnessetNum"], law["PrivateNumber"])
            scores[num][7] = "1"
            scores[num][8] = "0"
    else:
        if not str(law["StatusID"]) in [
            '106', '108', '109', '111', '113', '114',
            '115', '120', '122', '124', '126', '130',
            '140', '141', '141', '142', '143', '150',
            '161', '167', '169', '175', '178'
        ]:
            print("unknown StatusID", law["StatusID"],
                  law["Id"], law["KnessetNum"], law["PrivateNumber"])
        scores[num][7] = "1"
        scores[num][8] = "0"

    if DEBUG_STATUSES:
        continue

    old_names = None
    #print(doc["GroupTypeDesc"], doc["FilePath"])
    for p in get_doc(doc["FilePath"]).paragraphs:
        if init.match(p.text):  # .find("יוזמים:      חברי הכנסת") == 0:
            initiators = [a.strip() for a in split_initiators.split(
                init.sub("", p.text).replace("_", ""))]
            initiators = [a for a in initiators if a]
            laws_initiators[num] = initiators
            scores[num][9:] = initiators

        other_names = [i[0] for i in dups.findall(p.text)]
        if p.text.find("חוק") > 0 and other_names and (p.text.find("זהות") >= 0 or p.text.find("זהה") >= 0):
            old_names = other_names
            for n in old_names:
                for v in [n, n.replace('/', '\\'), n.replace('\\', '/')]:
                    if scored_laws.get(v):
                        name = v
                        if scored_laws[name]["שם הצעת החוק"].find(law['Name'][10:20]) == -1 and scored_laws[name]["שם הצעת החוק"].find(law['Name'][20:30]) == -1:
                            print(
                                "כפילות בשם שונה!", law['Name'], law_name, scored_laws[name]["שם הצעת החוק"], name)
                        if duplicates.get(law_name):
                            #print("law scored already twice! " + str(scored_laws[name]) + str(duplicates[law_name]))
                            if scored_laws[name]["ניקוד לחוק"] and duplicates[law_name]["ניקוד לחוק"] and scored_laws[name]["ניקוד לחוק"] != duplicates[law_name]["ניקוד לחוק"]:
                                print("חוק דורג פעמיים בעבר בניקוד שונה", law["Name"], scored_laws[name]["ניקוד לחוק"], scored_laws[
                                      name]["מספר חוק"], duplicates[law_name]["ניקוד לחוק"], duplicates[law_name]["מספר חוק"])
                        duplicates[law_name] = scored_laws[name]
                        old_csv += ('https://main.knesset.gov.il/Activity/Legislation/Laws/Pages/LawBill.aspx?t=lawsuggestionssearch&lawitemid=' + str(
                            law["Id"]) + ",\"" + law['Name'].replace("\"", "'") + "\"," + law_name + "," + scored_laws[name]["ניקוד לחוק"] + "," + scored_laws[name]["מספר חוק"] + "\n")
                        if not scores[int(law_name[5:])][3] and scored_laws[name]["ניקוד לחוק"] != "":
                            scores[int(law_name[5:])][1:6] = ["dup_laws_bot", law_name,
                                                              scored_laws[name]["ניקוד לחוק"], 'https://main.knesset.gov.il/Activity/Legislation/Laws/Pages/LawBill.aspx?t=lawsuggestionssearch&lawitemid=' + str(
                                                                  law["Id"]),
                                                              "\"" + (scored_laws[name]["הסבר הדירוג"] + " ראה חוק\n" + scored_laws[name]["מספר חוק"] + " " + scored_laws[name]["קישור להצעה"]).replace("\"", "'") + "\""]
                        break

    if not old_names:
        news_csv += ('https://main.knesset.gov.il/Activity/Legislation/Laws/Pages/LawBill.aspx?t=lawsuggestionssearch&lawitemid=' +
                     str(law["Id"]) + ",\"" + law['Name'].replace("\"", "'") + "\"," + law_name + "\n")
        if (scored_laws.get(name) == None or scored_laws.get(name) == "") and scores[num][3] == "":
            unscored_csv += ('https://main.knesset.gov.il/Activity/Legislation/Laws/Pages/LawBill.aspx?t=lawsuggestionssearch&lawitemid=' +
                             str(law["Id"]) + ",\"" + law['Name'].replace("\"", "'") + "\"," + law_name + "," + scores[num][7] + "\n")
    if not scores[num] or len(scores[num]) < 10 or not scores[num][9]:
        print("no iniitiators", [
              p.text for p in get_doc(doc["FilePath"]).paragraphs])

print(laws_last, len(scores), "laws")

if not DEBUG_STATUSES:
    open("unscored_laws.csv", "wt").write(unscored_csv)
    open("new_laws.csv", "wt").write(news_csv)
    open("scored_laws.csv", "wt").write(old_csv)
    open("table.csv", "wt").write(
        "\n".join([",".join([o or '' for o in line]) for line in scores if line]))
    open("initiators.csv", "wt").write(
        "\n".join([",".join([o or '' for o in line]) for line in laws_initiators]))

input()
